import os
import json
import google.generativeai as genai
from dotenv import load_dotenv
import requests
from PyPDF2 import PdfReader

# Carrega as variáveis do arquivo .env para o ambiente
load_dotenv()
# Configura a API do Google com a chave que está no .env
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
'''
Fase 3: Escrita (Assistente de Formatação) * Enquanto escreve seu próprio trabalho (em um editor de texto simples na plataforma ou externamente), o usuário pode usar o assistente de IA. * Ele pode colar um trecho de texto e pedir para a IA reescrevê-lo de forma mais acadêmica, ou, mais importante, gerar citações e referências no formato desejado (de acordo com a revista pra onde será publicado o artigo) a partir das informações de um artigo.
'''
'''
# Carrega as variáveis do arquivo .env para o ambiente
load_dotenv()

# Configura a API do Google com a chave que está no .env
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def extrair_texto_pdf(caminho_pdf: str) -> str:
    """
    Lê um arquivo PDF e retorna todo o texto concatenado.
    """
    texto_total = ""
    with open(caminho_pdf, "rb") as arquivo:
        leitor = PdfReader(arquivo)
        for pagina in leitor.pages:
            texto_total += pagina.extract_text() or ""
    return texto_total

def formatacao_academica(texto: str, estilo: str = "APA") -> str:
    """
    Usa a API do Google Generative AI para reescrever o texto em um estilo acadêmico específico.
    """
    prompt = f"Reescreva o seguinte texto em um estilo acadêmico seguindo o formato {estilo}:\n\n{texto}"
    
    resposta = genai.chat.completions.create(
        model="gemini-1.5-turbo",
        messages=[
            {"role": "system", "content": "Você é um assistente de formatação acadêmica."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000
    )
    
    return resposta.choices[0].message['content']'''

# formatacao_academica.py
import os
import json
from typing import Dict, Any, Optional
from dotenv import load_dotenv
import google.generativeai as genai

# Bibliotecas opcionais para PDF / DOCX
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

import requests

# Load env and configure Gemini (Google)
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)


# ---------------------------
# 1) Extrair texto de PDF
# ---------------------------
def extrair_texto_pdf(file_path: str) -> str:
    """
    Extrai texto de um arquivo PDF e retorna uma string com o texto agregado.
    Usa PyPDF2 se disponível. Lança exceção em erro.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"PDF não encontrado: {file_path}")

    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 não está instalado. Instale com `pip install PyPDF2`.")

    text_parts = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception as e:
                print(f"[warning] erro ao extrair página {i}: {e}")
                page_text = ""
            text_parts.append(page_text)
    full_text = "\n\n".join(text_parts).strip()
    print(f"[info] Extraído {len(text_parts)} páginas. Total de caracteres: {len(full_text)}")
    return full_text


# ---------------------------
# 2) Receber texto direto
# ---------------------------
def receber_texto_direto(raw_text: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    Normaliza/empacota texto recebido diretamente do front-end.
    Retorna dicionário com chaves: raw_text, title_hint, authors_hint, keywords_hint.
    """
    if metadata is None:
        metadata = {}

    payload = {
        "raw_text": raw_text.strip(),
        "metadata": {
            "title_hint": metadata.get("title_hint"),
            "authors_hint": metadata.get("authors_hint", []),
            "keywords_hint": metadata.get("keywords_hint", [])
        }
    }
    print(f"[info] Texto recebido (chars={len(payload['raw_text'])}), metadata keys: {list(payload['metadata'].keys())}")
    return payload


# ---------------------------
# 3) Prompt de formatação acadêmica (chamada ao Gemini)
# ---------------------------
PROMPT_TEMPLATE_SYSTEM = """
Você é um assistente especializado em formatar textos científicos para submissão a periódicos. 
Responda estritamente ao contrato JSON descrito abaixo — apenas JSON, sem explicações adicionais.
Siga estas regras globais:
1. Use o idioma: {language}.
2. Adapte o texto às regras do periódico {journal_name}.
3. Estilo de citação requisitado (CSL id): {csl_id}.
4. Limites: Abstract máximo {max_abstract_chars} caracteres.
5. Responda em conformidade com o schema JSON exigido.
6. Não invente referências. Se não houver DOI, marque resolução como UNRESOLVED.
"""

PROMPT_TEMPLATE_USER = """
Payload:
{{
  "journal_name": "{journal_name}",
  "csl_id": "{csl_id}",
  "language": "{language}",
  "max_abstract_chars": {max_abstract_chars},
  "title_case_rule": "{title_case_rule}",
  "required_sections": {required_sections},
  "raw_text": "<<RAW_TEXT>>",
  "metadata": {metadata_json}
}}

INSTRUÇÕES:
- Retorne apenas um objeto JSON seguindo o schema fornecido.
- Não inclua texto fora do JSON.
"""

# Schema string trimmed for brevity — the Gemini prompt expects the model to follow the schema we defined.
# (In production envie o schema completo ou referencie-o explicitamente.)
SCHEMA_BRIEF = """
OUTPUT_SCHEMA: O JSON deve conter: title, authors, abstract, keywords, sections, in_text_citations, references_raw, warnings, generation_info.
"""

def prompt_de_formatacao_academica(raw_text: str,
                                  journal_profile: Dict[str, Any],
                                  metadata: Optional[Dict[str, Any]] = None,
                                  model_name: str = "gemini-2.0-flash") -> Dict[str, Any]:
    """
    Monta o prompt com base no journal_profile e chama o Gemini para formatar academicamente o texto.
    Retorna o JSON parseado que o modelo deve entregar (conforme schema).
    """
    metadata = metadata or {}
    # Preencha placeholders
    system_msg = PROMPT_TEMPLATE_SYSTEM.format(
        language=journal_profile.get("language", "pt-BR"),
        journal_name=journal_profile.get("display_name", journal_profile.get("id", "unknown")),
        csl_id=journal_profile.get("citation_style", journal_profile.get("csl_id", "unknown")),
        max_abstract_chars=journal_profile.get("abstract", {}).get("max_chars", 250)
    )

    user_msg = PROMPT_TEMPLATE_USER.format(
        journal_name=journal_profile.get("display_name", journal_profile.get("id", "unknown")),
        csl_id=journal_profile.get("citation_style", journal_profile.get("csl_id", "unknown")),
        language=journal_profile.get("language", "pt-BR"),
        max_abstract_chars=journal_profile.get("abstract", {}).get("max_chars", 250),
        title_case_rule=journal_profile.get("title_case", "Sentence case"),
        required_sections=json.dumps(journal_profile.get("required_sections", []), ensure_ascii=False),
        metadata_json=json.dumps({
            "title_hint": metadata.get("title_hint"),
            "authors_hint": metadata.get("authors_hint", []),
            "keywords_hint": metadata.get("keywords_hint", [])
        }, ensure_ascii=False)
    )

    # Inject raw text (careful with size — chunk if necessary)
    # For simplicity, colocamos o raw_text direto; no ambiente real, você deve chunkar se grande.
    user_msg = user_msg.replace("<<RAW_TEXT>>", raw_text.replace("\n", "\\n"))

    prompt = f"{system_msg}\n{SCHEMA_BRIEF}\n{user_msg}\n\nPor favor retorne apenas o JSON."

    print("[info] Enviando prompt ao Gemini... (respeite limites de token)")

    if GOOGLE_API_KEY is None:
        raise RuntimeError("GOOGLE_API_KEY não configurado. Configure a variável de ambiente.")

    # Chamada ao Gemini (compatível com seu estilo original)
    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        text = response.text.strip()
        # Remove markers de codeblock se houver
        text = text.replace('```json', '').replace('```', '').strip()
        parsed = json.loads(text)
        print("[info] JSON recebido do Gemini e parseado com sucesso.")
        return parsed
    except json.JSONDecodeError as e:
        print(f"[error] JSON inválido recebido do Gemini: {e}. Conteúdo bruto segue:\n{text[:2000]}")
        raise
    except Exception as e:
        print(f"[error] Falha ao chamar Gemini: {e}")
        raise


# ---------------------------
# 4) Conversão para documento .docx
# ---------------------------
def conversao_texto_documento(formatted_json: Dict[str, Any],
                              output_path: str,
                              docx_template: Optional[str] = None) -> str:
    """
    Converte o JSON formatado (contrato do prompt) para um arquivo .docx.
    Estrutura simples:
      - Título
      - Autores / afiliações
      - Abstract
      - Seções (com headings)
      - Referências (simples)
    Tenta usar pypandoc/pandoc se disponível, caso contrário usa python-docx.
    Retorna o caminho do arquivo gerado.
    """
    # Primeiro: construa um markdown simples a partir do JSON
    title = formatted_json.get("title") or ""
    authors = formatted_json.get("authors", [])
    abstract = formatted_json.get("abstract") or ""
    sections = formatted_json.get("sections", [])
    references = formatted_json.get("references_raw", [])

    md_parts = []
    if title:
        md_parts.append(f"# {title}\n")
    if authors:
        author_lines = []
        for a in authors:
            name = a.get("name")
            aff = a.get("affiliation")
            if aff:
                author_lines.append(f"{name} — {aff}")
            else:
                author_lines.append(f"{name}")
        md_parts.append("\n".join(author_lines) + "\n")

    if abstract:
        md_parts.append("## Abstract\n")
        md_parts.append(abstract + "\n")

    for sec in sections:
        name = sec.get("name", "Section")
        content = sec.get("content", "")
        md_parts.append(f"## {name}\n")
        md_parts.append(content + "\n")

    if references:
        md_parts.append("## References\n")
        for r in references:
            raw = r.get("raw") if isinstance(r, dict) else str(r)
            md_parts.append(f"- {raw}")

    markdown = "\n\n".join(md_parts)

    # Tenta pandoc/pypandoc
    output_path = os.path.abspath(output_path)
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    if pypandoc is not None:
        try:
            extra_args = []
            if docx_template:
                extra_args = ["--reference-doc", docx_template]
            pypandoc.convert_text(markdown, to="docx", format="md", outputfile=output_path, extra_args=extra_args)
            print(f"[info] Documento gerado com pandoc em: {output_path}")
            return output_path
        except Exception as e:
            print(f"[warning] Falha ao gerar docx via pypandoc/pandoc: {e}. Tentando python-docx fallback...")

    # Fallback: python-docx (mais básico)
    if Document is None:
        raise RuntimeError("Nem pypandoc/pandoc nem python-docx estão disponíveis. Instale uma das bibliotecas.")

    doc = Document()
    # Title
    if title:
        h = doc.add_heading(title, level=0)
    # Authors
    if authors:
        p = doc.add_paragraph()
        p.add_run("\n".join([f"{a.get('name')}" + (f" — {a.get('affiliation')}" if a.get('affiliation') else "") for a in authors]))
    # Abstract
    if abstract:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(abstract)
    # Sections
    for sec in sections:
        doc.add_heading(sec.get("name", "Section"), level=1)
        doc.add_paragraph(sec.get("content", ""))
    # References
    if references:
        doc.add_heading("References", level=1)
        for r in references:
            raw = r.get("raw") if isinstance(r, dict) else str(r)
            doc.add_paragraph(raw, style='List Bullet')

    doc.save(output_path)
    print(f"[info] Documento gerado com python-docx em: {output_path}")
    return output_path


# ---------------------------
# Exemplo de uso (script)
# ---------------------------
if __name__ == "__main__":
    # Exemplo simples de journal_profile (substitua pelas suas configs reais)
    example_profile = {
        "id": "example_journal",
        "display_name": "Exemplo Journal",
        "citation_style": "apa",
        "language": "pt-BR",
        "abstract": {"max_chars": 200},
        "title_case": "Sentence case",
        "required_sections": ["Title", "Abstract", "Introduction", "Methods", "Results", "Conclusion", "References"]
    }

    # 1) Exemplo: extrair texto de PDF
    sample_pdf = "exemplo_artigo.pdf"
    if os.path.exists(sample_pdf) and PyPDF2 is not None:
        raw = extrair_texto_pdf(sample_pdf)
    else:
        # fallback: texto de demonstração
        raw = """Título: Exemplo de Estudo.
Resumo: Este é um resumo simples.
Introdução: Conteúdo introdutório.
Métodos: Descrevemos métodos.
Resultados: Observações.
Conclusão: Finalizamos."""

    # 2) Receber texto direto
    packaged = receber_texto_direto(raw, metadata={"title_hint": None, "authors_hint": [], "keywords_hint": []})

    # 3) Chamar Gemini para formatar (vai lançar se GOOGLE_API_KEY não configurada)
    try:
        formatted = prompt_de_formatacao_academica(packaged["raw_text"], example_profile, metadata=packaged["metadata"])
    except Exception as e:
        print(f"[error] Falha ao obter formatação pelo Gemini: {e}")
        # Em caso de falha, criamos um JSON mínimo para demonstração
        formatted = {
            "title": "Exemplo de Estudo",
            "authors": [{"name": "Autor Exemplo", "affiliation": None}],
            "abstract": "Este é um resumo simples.",
            "keywords": [],
            "sections": [
                {"name": "Introduction", "content": "Conteúdo introdutório."},
                {"name": "Methods", "content": "Descrevemos métodos."},
                {"name": "Results", "content": "Observações."},
                {"name": "Conclusion", "content": "Finalizamos."}
            ],
            "in_text_citations": [],
            "references_raw": [],
            "warnings": [],
            "generation_info": {"model": "local-fallback", "prompt_template_version": "v1"}
        }

    # 4) Converter para docx
    out_file = os.path.join(os.getcwd(), "output_example.docx")
    try:
        path = conversao_texto_documento(formatted, out_file)
        print(f"[success] Arquivo final disponível em: {path}")
    except Exception as e:
        print(f"[error] Falha ao converter para documento: {e}")
